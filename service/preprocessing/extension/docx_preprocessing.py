"""
DOCX 전처리 모듈
DOCX 파일을 텍스트와 표로 추출하는 기능 제공
- 표 헤더 감지 및 구조화
- 스타일 정보 활용 (제목, 본문 등)
- 하이퍼링크 추출
- 리스트 처리 (번호 목록, 불릿 목록)
- 강조 텍스트 처리 (볼드, 이탤릭)
"""
from __future__ import annotations
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def _clean_text(s: str | None) -> str:
    """텍스트 정규화 (공통 유틸리티에서 import)"""
    from service.preprocessing.extension.utils import clean_text as clean_text_util
    return clean_text_util(s)


def _extract_hyperlink(paragraph) -> str:
    """단락에서 하이퍼링크 추출"""
    text_parts = []
    try:
        for run in paragraph.runs:
            try:
                if hasattr(run, 'hyperlink') and run.hyperlink:
                    link_text = run.text.strip() if run.text else ""
                    link_url = run.hyperlink.target if hasattr(run.hyperlink, 'target') else ""
                    if link_text:
                        if link_url:
                            text_parts.append(f"{link_text} ({link_url})")
                        else:
                            text_parts.append(link_text)
                else:
                    if run.text:
                        text_parts.append(run.text)
            except Exception as e:
                logger.debug(f"Failed to extract hyperlink from run: {e}")
                if run.text:
                    text_parts.append(run.text)
    except Exception as e:
        logger.warning(f"Failed to extract hyperlink from paragraph: {e}")
        return paragraph.text if paragraph.text else ""
    return "".join(text_parts)


def _extract_formatted_text(paragraph) -> str:
    """단락에서 포맷팅된 텍스트 추출 (볼드, 이탤릭 등)"""
    text_parts = []
    try:
        for run in paragraph.runs:
            try:
                text = run.text if run.text else ""
                if not text:
                    continue
                
                # 볼드 처리
                if hasattr(run, 'bold') and run.bold:
                    text = f"**{text}**"
                # 이탤릭 처리
                if hasattr(run, 'italic') and run.italic:
                    text = f"*{text}*"
                
                text_parts.append(text)
            except Exception as e:
                logger.debug(f"Failed to extract formatted text from run: {e}")
                if run.text:
                    text_parts.append(run.text)
    except Exception as e:
        logger.warning(f"Failed to extract formatted text from paragraph: {e}")
        return paragraph.text if paragraph.text else ""
    return "".join(text_parts)


def _table_to_markdown(table) -> str:
    """표를 마크다운 형식으로 변환 (헤더 감지 및 구조화)"""
    if not table.rows:
        return ""
    
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = _clean_text(cell.text).strip()
            # 빈 셀은 공백으로 처리
            if not cell_text:
                cell_text = ""
            # 셀 내부의 줄바꿈을 공백으로 변환 (마크다운 테이블 호환)
            cell_text = cell_text.replace("\n", " ").replace("\r", " ")
            cells.append(cell_text)
        rows.append(cells)
    
    if not rows:
        return ""
    
    # 헤더 행 감지: 첫 번째 행이 비어있지 않고, 두 번째 행과 구조가 유사하면 헤더로 간주
    # 또는 첫 번째 행의 모든 셀이 비어있지 않으면 헤더로 간주
    has_header = False
    if len(rows) > 1:
        first_row = rows[0]
        # 첫 번째 행의 모든 셀이 비어있지 않고, 텍스트가 짧고 명확하면 헤더로 간주
        if all(cell.strip() for cell in first_row) and len(first_row) > 0:
            # 첫 번째 행의 평균 길이가 짧으면 헤더일 가능성 높음
            avg_len = sum(len(cell) for cell in first_row) / len(first_row) if first_row else 0
            if avg_len < 50:  # 평균 50자 미만이면 헤더로 간주
                has_header = True
    
    # 마크다운 테이블 생성
    lines = []
    
    # 헤더 행
    if has_header and rows:
        header_row = rows[0]
        lines.append("| " + " | ".join(header_row) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_row)) + " |")
        data_rows = rows[1:]
    else:
        # 헤더가 없으면 첫 번째 행도 데이터로 처리
        data_rows = rows
    
    # 데이터 행
    for row in data_rows:
        # 행의 셀 수가 헤더와 다를 수 있으므로 맞춤
        if has_header and len(row) != len(rows[0]):
            # 셀 수가 다르면 헤더 수에 맞춤
            row = row[:len(rows[0])] + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(row) + " |")
    
    return "\n".join(lines)


def _convert_doc_to_docx(src: Path) -> Path | None:
    """DOC 파일을 DOCX로 변환 (LibreOffice 사용)"""
    try:
        import subprocess
        import shutil
        from pathlib import Path
        
        # LibreOffice 명령어 확인
        libreoffice_cmd = shutil.which("libreoffice")
        if not libreoffice_cmd:
            logger.warning(f"[DOC->DOCX] LibreOffice가 설치되지 않았습니다. DOC 파일 변환을 건너뜁니다: {src.name}")
            return None
        
        outdir = src.parent
        logger.info(f"[DOC->DOCX] 변환 시작: {src.name} -> {outdir}")
        
        # stdout과 stderr를 캡처하여 디버깅 정보 확인
        result = subprocess.run(
            [
                libreoffice_cmd, "--headless", "--convert-to", "docx",
                "--outdir", str(outdir), str(src)
            ],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
        
        # stdout과 stderr 로그 출력
        if result.stdout:
            stdout_msg = result.stdout.decode("utf-8", errors="ignore")
            if stdout_msg.strip():
                logger.debug(f"[DOC->DOCX] stdout: {stdout_msg[:500]}")
        
        if result.returncode != 0:
            stderr_msg = result.stderr.decode("utf-8", errors="ignore") if result.stderr else ""
            logger.error(
                f"[DOC->DOCX] 변환 실패 (returncode={result.returncode}): {src.name}. "
                f"에러: {stderr_msg[:500]}"
            )
            return None
        else:
            logger.info(f"[DOC->DOCX] LibreOffice 변환 명령 완료 (returncode=0): {src.name}")
        
        # 변환된 파일 찾기 (LibreOffice는 원본 파일명을 유지하되 확장자만 변경)
        cand = src.with_suffix(".docx")
        
        # 파일이 생성될 때까지 잠시 대기 (비동기 처리 대비)
        import time
        max_wait = 10  # 최대 10초 대기 (증가)
        wait_interval = 0.5  # 0.5초마다 확인
        waited = 0
        logger.debug(f"[DOC->DOCX] 변환된 파일 대기 중: {cand}")
        while not cand.exists() and waited < max_wait:
            time.sleep(wait_interval)
            waited += wait_interval
            if waited % 2 == 0:  # 2초마다 로그
                logger.debug(f"[DOC->DOCX] 파일 대기 중... ({waited:.1f}초)")
        
        if cand.exists():
            file_size = cand.stat().st_size
            logger.info(f"[DOC->DOCX] 변환 성공: {src.name} -> {cand.name} (크기: {file_size} bytes)")
            return cand
        else:
            # 다른 가능한 경로 확인 (대소문자 차이 등)
            parent = src.parent
            stem = src.stem
            possible_names = [
                f"{stem}.docx",
                f"{stem}.DOCX",
                f"{stem}.Docx",
            ]
            logger.debug(f"[DOC->DOCX] 대체 경로 검색 중: {possible_names}")
            for name in possible_names:
                alt_path = parent / name
                if alt_path.exists():
                    file_size = alt_path.stat().st_size
                    logger.info(f"[DOC->DOCX] 변환 성공 (대체 경로): {src.name} -> {alt_path.name} (크기: {file_size} bytes)")
                    return alt_path
            
            # 디렉토리 전체 내용 확인
            all_files = list(parent.glob(f'{stem}.*'))
            logger.error(f"[DOC->DOCX] 변환된 파일을 찾을 수 없습니다. 예상 경로: {cand}")
            logger.error(f"[DOC->DOCX] 디렉토리 내용: {[f.name for f in all_files]}")
            logger.error(f"[DOC->DOCX] 원본 파일 존재: {src.exists()}, 경로: {src}")
            return None
            
    except subprocess.TimeoutExpired:
        logger.error(f"[DOC->DOCX] 변환 타임아웃: {src.name} (120초 초과)")
        return None
    except FileNotFoundError:
        logger.error(f"[DOC->DOCX] LibreOffice 명령어를 찾을 수 없습니다. 설치가 필요합니다.")
        return None
    except Exception as e:
        logger.exception(f"[DOC->DOCX] 변환 중 예외 발생: {src.name}, 오류: {e}")
        return None


def extract_docx(fp: Path) -> tuple[str, list[dict]]:
    """DOCX 파일 추출 (고도화된 버전)
    DOC 파일이 들어오면 자동으로 DOCX로 변환 후 처리합니다.
    
    Returns:
        tuple[str, list[dict]]: (본문 텍스트, 표 리스트)
    """
    # DOC 파일인 경우 DOCX로 변환
    converted_path = None
    is_converted = False
    original_path = fp
    
    if fp.suffix.lower() == ".doc":
        logger.info(f"[DOCX] DOC 파일 감지, DOCX로 변환 시작: {fp.name} (경로: {fp})")
        if not fp.exists():
            logger.error(f"[DOCX] 원본 DOC 파일이 존재하지 않습니다: {fp}")
            return "", []
        
        converted_path = _convert_doc_to_docx(fp)
        if converted_path and converted_path.exists():
            file_size = converted_path.stat().st_size
            logger.info(f"[DOCX] 변환 성공, 변환된 파일 사용: {converted_path.name} (크기: {file_size} bytes)")
            fp = converted_path
            is_converted = True
        else:
            logger.error(f"[DOCX] DOC 파일 변환 실패: {original_path.name}. 변환된 파일이 없습니다.")
            logger.error(f"[DOCX] 원본 파일 경로: {original_path}, 존재 여부: {original_path.exists()}")
            return "", []
    
    try:
        from docx import Document
    except ImportError:
        logger.warning(f"python-docx not available, treating {fp} as plain text")
        # 변환된 파일이 있으면 삭제
        if is_converted and converted_path and converted_path.exists():
            try:
                converted_path.unlink()
            except Exception:
                pass
        from service.preprocessing.extension.txt_preprocessing import _extract_plain_text
        return _extract_plain_text(fp)
    
    try:
        # DOCX 파일 열기 (경로를 문자열로 변환하여 전달)
        doc_path = str(fp.resolve())
        d = Document(doc_path)
        text_parts = []
        tables = []
        
        # 표 추출
        for tb in d.tables:
            try:
                md_table = _table_to_markdown(tb)
                if md_table:
                    tables.append({
                        "page": 0,  # DOCX는 페이지 정보가 없음
                        "bbox": [],
                        "text": md_table
                    })
            except Exception as e:
                logger.warning(f"Failed to extract table from DOCX {fp}: {e}")
                continue
        
        # 단락 추출 (순서대로 처리)
        for para in d.paragraphs:
            try:
                # para.text는 이미 UTF-8로 디코딩된 문자열을 반환합니다
                para_text = para.text.strip()
                
                if not para_text:
                    continue
                
                # 스타일 정보 확인
                style_name = para.style.name if para.style else ""
                
                # 제목 스타일 처리
                if style_name.startswith('Heading') or style_name.startswith('제목') or 'heading' in style_name.lower():
                    level = 1
                    # Heading 1, Heading 2 등에서 레벨 추출
                    if 'Heading' in style_name or 'heading' in style_name.lower():
                        try:
                            # "Heading 1", "Heading 2" 등에서 숫자 추출
                            parts = style_name.split()
                            for part in parts:
                                if part.isdigit():
                                    level = int(part)
                                    break
                        except (ValueError, IndexError):
                            level = 1
                    
                    # 마크다운 제목 형식으로 변환
                    heading_prefix = "#" * min(level, 6)  # 최대 6레벨
                    para_text = f"{heading_prefix} {para_text}"
                
                # 리스트 처리 (리스트 스타일이 있으면 불릿 또는 번호 추가)
                elif para.style and ('List' in para.style.name or 'list' in para.style.name.lower()):
                    # 리스트는 그대로 유지하되, 앞에 불릿 추가 (필요시)
                    if not para_text.startswith(('•', '-', '*', '1.', '2.', '3.')):
                        para_text = f"• {para_text}"
                else:
                    # 일반 단락: 하이퍼링크 및 포맷팅 추출 시도
                    try:
                        # 하이퍼링크가 있는지 확인
                        has_hyperlink = any(
                            hasattr(run, 'hyperlink') and run.hyperlink 
                            for run in para.runs
                        )
                        
                        # 포맷팅이 있는지 확인
                        has_formatting = any(
                            (hasattr(run, 'bold') and run.bold) or 
                            (hasattr(run, 'italic') and run.italic)
                            for run in para.runs
                        )
                        
                        if has_hyperlink:
                            para_text = _extract_hyperlink(para)
                        elif has_formatting:
                            para_text = _extract_formatted_text(para)
                    except Exception as e:
                        # 추출 실패 시 기본 텍스트 사용
                        logger.debug(f"Failed to extract formatting from paragraph: {e}")
                        pass
                
                if para_text:
                    # UTF-8 인코딩 확인 및 정리
                    cleaned = _clean_text(para_text)
                    if cleaned:
                        text_parts.append(cleaned)
            except Exception as e:
                logger.warning(f"Failed to extract paragraph from DOCX {fp}: {e}")
                continue
        
        main_text = "\n\n".join(text_parts)
        cleaned_text = _clean_text(main_text)
        
        # 한글이 제대로 추출되었는지 확인
        if not cleaned_text or len(cleaned_text.encode('utf-8')) < 10:
            logger.warning(f"DOCX extraction may have failed for {fp}, extracted text is too short or empty")
        
        return cleaned_text, tables
        
    except Exception as e:
        logger.exception(f"Failed to extract DOCX {fp}: {e}")
        # DOCX 파일은 바이너리이므로 일반 텍스트로 읽으면 안 됨
        # 폴백하지 않고 빈 결과 반환
        return "", []
    finally:
        # 변환된 임시 DOCX 파일 삭제
        if is_converted and converted_path and converted_path.exists():
            try:
                converted_path.unlink()
                logger.debug(f"[DOCX] 임시 변환 파일 삭제: {converted_path}")
            except Exception as e:
                logger.warning(f"[DOCX] 임시 변환 파일 삭제 실패: {converted_path}, 오류: {e}")

